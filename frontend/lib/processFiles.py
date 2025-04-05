import os
import re
import json
from docx import Document
import google.generativeai as genai
import pandas as pd
import tempfile
from pdf2image import convert_from_path
from PIL import Image

# Configure Gemini API
genai.configure(api_key="AIzaSyAQvW-7i3jnNu5qwolDOPV9q2HhdkKtrAU")

def extract_questions_from_docx(docx_path):
    """Extract questions from a DOCX file."""
    doc = Document(docx_path)
    questions = []
    pattern = re.compile(r"Question\s+\d+[:.]?\s*(.*?)\s*\[(\d+)\s*marks?\]", re.IGNORECASE | re.DOTALL)
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        match = pattern.search(text)
        if match:
            question, marks = match.groups()
            questions.append((question.strip(), int(marks)))
    
    return questions

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file using Gemini Vision."""
    model = genai.GenerativeModel("gemini-1.5-flash")
    prompt = """Extract all text from this document exactly as it appears. 
    Pay special attention to questions in formats like:
    - Question 1 [5 marks]
    - Question 2: Describe... [10 marks]
    - Q3. Explain... [3 marks]"""
    
    images = convert_from_path(pdf_path, dpi=300)
    extracted_text = ""
    
    for image in images:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as temp_img:
            image.save(temp_img.name, format="JPEG")
        
        with Image.open(temp_img.name) as img:
            response = model.generate_content([prompt, img], stream=False)
            extracted_text += response.text + "\n"
        
        os.unlink(temp_img.name)
    
    return extracted_text

def extract_questions_from_pdf(pdf_path):
    """Extract questions from a PDF file."""
    questions = []
    pattern = re.compile(
        r"(?:Question|Q)\s*\d+[:.]?\s*(.*?)\s*\[(\d+)\s*marks?\]", 
        re.IGNORECASE | re.DOTALL
    )
    
    extracted_text = extract_text_from_pdf(pdf_path)
    matches = pattern.finditer(extracted_text)
    
    for match in matches:
        question = match.group(1).strip()
        marks = int(match.group(2))
        questions.append((question, marks))
    
    return questions

def classify_questions(questions):
    """Classify questions using Gemini."""
    classified_data = []
    model = genai.GenerativeModel("gemini-1.5-flash")
    
    for i, (question, marks) in enumerate(questions):
        prompt = f"""
        Analyze the following question and classify it as exactly one of these types:
        - "MCQ" (if it's a multiple choice question)
        - "Descriptive" (if it requires a written explanation)
        - "Opinion-based" (if it asks for personal opinion)
        
        Question: {question}
        
        Return ONLY the classification word (MCQ, Descriptive, or Opinion-based) with no other text.
        """
        
        try:
            response = model.generate_content(prompt)
            predicted_type = response.text.strip()
            
            # Validate response
            if predicted_type not in ["MCQ", "Descriptive", "Opinion-based"]:
                # If invalid, make a more specific request
                prompt += "\n\nPlease respond with ONLY one word: MCQ, Descriptive, or Opinion-based"
                response = model.generate_content(prompt)
                predicted_type = response.text.strip()
                
                # Default to Descriptive if still invalid
                if predicted_type not in ["MCQ", "Descriptive", "Opinion-based"]:
                    predicted_type = "Descriptive"
        except Exception as e:
            print(f"Error classifying question: {e}")
            predicted_type = "Descriptive"
        
        classified_data.append({
            "Question Number": i + 1,
            "Question": question,
            "Type": predicted_type,
            "Marks": marks
        })
    
    return classified_data

def save_questions_to_json(classified_data, output_json):
    """Save classified questions to a JSON file."""
    os.makedirs(os.path.dirname(output_json), exist_ok=True)
    with open(output_json, "w", encoding="utf-8") as f:
        json.dump(classified_data, f, indent=4, ensure_ascii=False)

def get_answer(question, marks):
    """Generate an answer using Gemini API."""
    model = genai.GenerativeModel("gemini-1.5-flash")
    
    if marks == 1:
        prompt = f"Answer this in one precise sentence:\nQ: {question}"
    elif marks == 2:
        prompt = f"Provide an answer in 2-3 sentences:\nQ: {question}"
    elif marks == 3:
        prompt = f"Provide a short answer in 3-4 sentences:\nQ: {question}"
    elif marks == 4:
        prompt = f"Provide a detailed answer in about 5-6 sentences:\nQ: {question}"
    elif marks >= 5:
        prompt = f"Write a comprehensive answer in multiple paragraphs (about 500 words):\nQ: {question}"
    else:
        prompt = f"Provide a precise and clear answer:\nQ: {question}"

    try:
        response = model.generate_content(prompt)
        return response.text.strip() if response.text else "No answer generated."
    except Exception as e:
        print(f"Error generating answer: {e}")
        return "Error generating answer."

def generate_answers_json_and_docx(classified_data, output_json, output_docx):
    """Generate answers and save them in JSON and DOCX formats."""
    # First save the classified data without answers
    save_questions_to_json(classified_data, output_json)
    
    # Now generate answers and update the data
    for entry in classified_data:
        entry["Reference Answer"] = get_answer(entry["Question"], entry["Marks"])
    
    # Save the updated data with answers
    save_questions_to_json(classified_data, output_json)
    
    # Generate DOCX file
    doc = Document()
    doc.add_heading('Reference Answers', level=1)
    
    for entry in classified_data:
        doc.add_heading(f"Question {entry['Question Number']}", level=2)
        doc.add_paragraph(entry["Question"])
        doc.add_paragraph(f"Marks: {entry['Marks']}")
        doc.add_paragraph("Reference Answer:")
        doc.add_paragraph(entry["Reference Answer"])
        doc.add_paragraph("\n")
    
    doc.save(output_docx)

def process_question_paper(input_file):
    """Process the question paper and generate reference answers."""
    output_dir = "../frontend/outputs"
    os.makedirs(output_dir, exist_ok=True)
    output_json = os.path.join(output_dir, "classified_questions.json")
    output_docx = os.path.join(output_dir, "answers.docx")
    
    print(f"Processing {input_file}...")
    
    try:
        if input_file.endswith(".pdf"):
            print("Extracting questions from PDF...")
            questions = extract_questions_from_pdf(input_file)
        elif input_file.endswith(".docx"):
            print("Extracting questions from DOCX...")
            questions = extract_questions_from_docx(input_file)
        else:
            raise ValueError("Unsupported file format. Only PDF and DOCX are supported.")
        
        print(f"Found {len(questions)} questions. Classifying them...")
        classified_data = classify_questions(questions)
        
        print("Generating reference answers...")
        generate_answers_json_and_docx(classified_data, output_json, output_docx)
        
        print(f"Processing complete. Files saved to:\n- {output_json}\n- {output_docx}")
        return True
    except Exception as e:
        print(f"Error processing question paper: {e}")
        return False

def extract_text_from_image(image):
    """Extract text from an image using Gemini 1.5 Flash."""
    model = genai.GenerativeModel("gemini-1.5-flash")
    prompt = "Extract only the text from this image. Do not add any explanations or extra words. Maintain the exact structure of the text."
    response = model.generate_content([image, prompt], stream=False)
    return response.text.strip()

def extract_answers_from_pdf(pdf_path):
    """Extract answers from a student's PDF using Gemini OCR."""
    answers = []
    roll_number = "Unknown"

    # Convert PDF to images
    images = convert_from_path(pdf_path, dpi=300)

    for page_num, image in enumerate(images):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as temp_img:
            image.save(temp_img.name, format="JPEG")
        
        # Open the image file and ensure it is closed after processing
        with Image.open(temp_img.name) as img:
            extracted_text = extract_text_from_image(img)
        
        # Delete the temporary file after processing
        os.unlink(temp_img.name)

        # Debugging: Print extracted text
        print(f"Page {page_num + 1} Extracted Text:\n{extracted_text}\n")

        if extracted_text:
            if page_num == 0:  # First page might contain Roll Number
                # Look for roll number in the extracted text
                if "Roll Number" in extracted_text:
                    roll_line = extracted_text.split("Roll Number")[1].split("\n")[0].strip()
                    roll_number = roll_line.split()[-1].strip()  # Extract the roll number

            # Split text into lines
            lines = extracted_text.split("\n")
            current_answer = []  # Temporary list to hold lines for the current answer
            for line in lines:
                line = line.strip()
                if not line:
                    continue  # Skip empty lines

                # Check if the line starts with "Ans" (case-insensitive)
                if line.lower().startswith("ans"):
                    # If we have collected lines for a previous answer, add it to the answers list
                    if current_answer:
                        answers.append(" ".join(current_answer).strip())  # Join lines into a single answer
                        current_answer = []  # Reset for the next answer

                    # Add the current "Ans" line to the new answer
                    current_answer.append(line)
                else:
                    # If the line doesn't start with "Ans", it's part of the current answer
                    current_answer.append(line)

            # Add the last collected answer (if any)
            if current_answer:
                answers.append(" ".join(current_answer).strip())

    return answers, roll_number

def evaluate_answer(question, reference_answer, student_answer, q_type, marks):
    """Use Gemini to evaluate student answers."""
    prompt = f"""
    You are an expert teacher evaluating a student's answer.
    
    Question: {question}
    Question Type: {q_type}
    Reference Answer: {reference_answer}
    Student Answer: {student_answer}
    Maximum Marks: {marks}
    
    Based on comparison between the reference and student answer, assign a score out of {marks}.
    For MCQs, give full marks only for exact matches.
    For Descriptive answers, evaluate based on factual accuracy, completeness, and clarity.
    For Opinion-based answers, evaluate based on coherence, logical flow, and relevance.
    
    Return only the numeric score as a float (with up to 1 decimal place). Do not include any explanation or additional text.
    """
    
    model = genai.GenerativeModel("gemini-1.5-flash-latest")
    response = model.generate_content(prompt)
    
    try:
        score = float(response.text.strip())
        return min(score, marks)  # Ensure score doesn't exceed max marks
    except ValueError:
        # Fallback if Gemini doesn't return a valid score
        if q_type == "MCQ":
            return marks if reference_answer.lower() == student_answer.lower() else 0.0
        else:
            # Default to half marks if parsing fails
            return marks / 2

def identify_weak_topics(question, student_answer, reference_answer):
    """Use Gemini to identify weak topics and improvement tips."""
    prompt = f"""
    Analyze the student's mistakes in answering the following question. Identify the weak topic and suggest improvements.
    
    Question: {question}
    Reference Answer: {reference_answer}
    Student Answer: {student_answer}
    
    Provide the weak topic and tips in JSON format with keys 'weak_topic' and 'tips'. The 'tips' should be a list of strings.
    Example: {{"weak_topic": "Topic Name", "tips": ["Tip 1", "Tip 2", "Tip 3"]}}
    """
    
    model = genai.GenerativeModel("gemini-1.5-flash-latest")
    response = model.generate_content(prompt)
    
    try:
        result = json.loads(response.text)
        return result
    except json.JSONDecodeError:
        # Fallback if JSON parsing fails
        weak_topic_match = re.search(r'"weak_topic"\s*:\s*"([^"]+)"', response.text)
        tips_match = re.findall(r'"tips"\s*:\s*\[([^\]]+)\]', response.text)
        
        weak_topic = weak_topic_match.group(1) if weak_topic_match else "Unknown"
        tips = []
        if tips_match:
            tips = [tip.strip().strip('"\'') for tip in tips_match[0].split(",")]
        
        return {"weak_topic": weak_topic, "tips": tips or ["No specific tips available."]}

def evaluate_answers(input_pdf, classified_json, output_json):
    """Evaluate a student's answers against reference answers."""
    try:
        with open(classified_json, "r", encoding="utf-8") as f:
            classified_data = json.load(f)
        
        student_answers, roll_number = extract_answers_from_pdf(input_pdf)
        
        results = []
        
        for i, entry in enumerate(classified_data):
            question_num = entry["Question Number"]
            question = entry["Question"]
            q_type = entry["Type"]
            marks = entry["Marks"]
            reference_answer = entry["Reference Answer"]
            
            student_answer = student_answers[i] if i < len(student_answers) else ""
            
            score = evaluate_answer(question, reference_answer, student_answer, q_type, marks)
            
            result_entry = {
                "roll_number": roll_number,
                "file_name": os.path.basename(input_pdf),
                "question_number": question_num,
                "question": question,
                "reference_answer": reference_answer,
                "student_answer": student_answer,
                "marks_awarded": score,
                "total_marks": marks
            }
            
            if score < marks * 0.8: 
                weak_data = identify_weak_topics(question, student_answer, reference_answer)
                result_entry["weak_topics"] = weak_data["weak_topic"]
                result_entry["improvement_tips"] = weak_data["tips"]
            
            results.append(result_entry)
        
        os.makedirs(os.path.dirname(output_json), exist_ok=True)
        
        with open(output_json, "w", encoding="utf-8") as f:
            json.dump(results, f, indent=4, ensure_ascii=False)
        
        return results
    
    except Exception as e:
        print(f"An error occurred while evaluating answers: {e}")
        return []

def process_student_answer_sheet(student_pdf):
    """Process a single student's answer sheet."""
    classified_json = "../frontend/outputs/classified_questions.json"
    output_json = f"../frontend/data/{os.path.basename(student_pdf).split('.')[0]}_results.json"
    
    results = evaluate_answers(student_pdf, classified_json, output_json)
    print(f"Evaluation complete for {student_pdf}. Results saved to {output_json}")
    
    return output_json

def process_all_student_sheets(folder_path):
    """Process all student answer sheets in a folder."""
    result_files = []
    
    for filename in os.listdir(folder_path):
        if filename.endswith(".pdf") and "question-paper" not in filename.lower():
            student_pdf = os.path.join(folder_path, filename)
            result_file = process_student_answer_sheet(student_pdf)
            result_files.append(result_file)
    
    return result_files

def generate_class_summary(result_files, output_json):
    """Generate a summary of all students' performance."""
    all_results = []
    student_summaries = {}
    
    for file_path in result_files:
        with open(file_path, "r", encoding="utf-8") as f:
            results = json.load(f)
            all_results.extend(results)
            
            if results:
                roll_number = results[0]["roll_number"]
                total_marks = sum(r["total_marks"] for r in results)
                marks_awarded = sum(r["marks_awarded"] for r in results)
                percentage = (marks_awarded / total_marks * 100) if total_marks > 0 else 0
                
                weak_topics = {}
                for result in results:
                    if "weak_topics" in result:
                        topic = result["weak_topics"]
                        weak_topics[topic] = weak_topics.get(topic, 0) + 1
                
                common_weak_topics = [topic for topic, count in 
                                     sorted(weak_topics.items(), key=lambda x: x[1], reverse=True)[:3]]
                
                student_summaries[roll_number] = {
                    "roll_number": roll_number,
                    "total_marks": total_marks,
                    "marks_awarded": marks_awarded,
                    "percentage": round(percentage, 2),
                    "common_weak_topics": common_weak_topics
                }
    
    summary = {
        "individual_results": all_results,
        "student_summaries": list(student_summaries.values())
    }
    
    with open(output_json, "w", encoding="utf-8") as f:
        json.dump(summary, f, indent=4, ensure_ascii=False)
    
    print(f"Class summary generated and saved to {output_json}")

if __name__ == "__main__":
    folder = "../frontend/uploads"
    base_filename = "question-paper"
    
    pdf_file = os.path.join(folder, base_filename + ".pdf")
    docx_file = os.path.join(folder, base_filename + ".docx")

    if os.path.exists(pdf_file):
        process_question_paper(pdf_file)
    elif os.path.exists(docx_file):
        process_question_paper(docx_file)
    else:
        print("No valid question paper file found!")
        exit(1)
    
    result_files = process_all_student_sheets(folder)
    folderr="../frontend/data"
    summary_json = os.path.join(folderr,"class_summary.json")
    generate_class_summary(result_files, summary_json)
    
    print("All processing completed successfully!")